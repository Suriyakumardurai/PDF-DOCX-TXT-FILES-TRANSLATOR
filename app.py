import os
from flask import Flask, request, render_template, send_file
from googletrans import Translator
import fitz
from docx import Document

app = Flask(__name__)

UPLOAD_FOLDER = 'uploads'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

def translate_text(text, target_language='en'):
    translator = Translator()
    translated_text = translator.translate(text, dest=target_language)
    return translated_text.text

def translate_file(input_file, target_language='en'):
    translated_content = ""
    output_file_path = ""

    if input_file.lower().endswith('.txt'):
        with open(input_file, 'r', encoding='utf-8') as file:
            content = file.read()
            translated_content = translate_text(content, target_language)
            output_extension = '.txt'

    elif input_file.lower().endswith('.pdf'):
        pdf_document = fitz.open(input_file)
        translated_content = ""
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            translated_content += page.get_text()

        translated_content = translate_text(translated_content, target_language)
        output_extension = '.pdf'

    elif input_file.lower().endswith('.docx'):
        doc = Document(input_file)
        for paragraph in doc.paragraphs:
            translated_content += paragraph.text + '\n'
        translated_content = translate_text(translated_content, target_language)
        output_extension = '.docx'

    else:
        print("Unsupported file format.")
        return None

    output_file_path = os.path.join(app.config['UPLOAD_FOLDER'], os.path.splitext(os.path.basename(input_file))[0] + "_translated" + output_extension)

    if output_extension == '.txt':
        with open(output_file_path, 'w', encoding='utf-8') as output_file:
            output_file.write(translated_content)
    elif output_extension == '.pdf':
        pdf_document = fitz.open()
        pdf_page = pdf_document.new_page()
        pdf_page.insert_text((50, 50), translated_content, fontname="helv", fontsize=12)
        pdf_document.save(output_file_path)
    elif output_extension == '.docx':
        doc = Document()
        doc.add_paragraph(translated_content)
        doc.save(output_file_path)

    return output_file_path

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        target_language = request.form['target_language']
        uploaded_file = request.files['file']
        if uploaded_file.filename != '':
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], uploaded_file.filename)
            uploaded_file.save(file_path)
            translated_file = translate_file(file_path, target_language)
            if translated_file:
                return send_file(translated_file, as_attachment=True, download_name=os.path.basename(translated_file))
            else:
                return "Translation failed. Unsupported file format."
    return render_template('index.html')

if __name__ == '__main__':
    app.run(debug=True,host="0.0.0.0", port= 8000)
